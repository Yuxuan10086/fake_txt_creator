#!/usr/bin/python
# -*- coding: UTF-8 -*-

import os, re
import random,readJSON

data = readJSON.读JSON文件("data.json")
名人名言 = data["famous"] # a 代表前面垫话，b代表后面垫话
前面垫话 = data["before"] # 在名人名言前面弄点废话
后面垫话 = data['after']  # 在名人名言后面弄点废话
废话 = data['bosh'] # 代表文章主要废话来源

xx = "学生会退会"

重复度 = 2

def 洗牌遍历(列表):
    global 重复度
    池 = list(列表) * 重复度
    while True:
        random.shuffle(池)
        for 元素 in 池:
            yield 元素

下一句废话 = 洗牌遍历(废话)
下一句名人名言 = 洗牌遍历(名人名言)

def 来点名人名言():
    global 下一句名人名言
    xx = next(下一句名人名言)
    xx = xx.replace(  "a",random.choice(前面垫话) )
    xx = xx.replace(  "b",random.choice(后面垫话) )
    return xx

def 另起一段():
    xx = ". "
    xx += "\r\n"
    xx += "    "
    return xx

def creator(xx):
    for x in xx:
        tmp = str()
        while ( len(tmp) < 6000 ) :
            分支 = random.randint(0,100)
            if 分支 < 5:
                tmp += 另起一段()
            elif 分支 < 20 :
                tmp += 来点名人名言()
            else:
                tmp += next(下一句废话)
        tmp = tmp.replace("x",xx)
        return tmp

def main():
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    import csv
    f = open(os.path.split(os.path.realpath(__file__))[0] + '\\name.csv', 'r', encoding='utf-8')
    with f:
        reader = csv.reader(f)
        for row in reader:
            for e in row:
                font = "宋体"
                doc = Document()
                para = doc.add_paragraph(" ")
                for i in creator("核武器世界和平"):
                    run = para.add_run(i)
                    run.font.name = font
                    run.font.size = 300000
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
                doc.save(os.path.split(os.path.realpath(__file__))[0] + "\\res\\" + e + ".docx")

main()